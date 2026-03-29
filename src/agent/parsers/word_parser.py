"""Word parser for .docx files.

This parser handles Word documents (.docx), extracting:
- Paragraphs and text content
- Headings (by style)
- Tables
- Document structure
"""

from dataclasses import dataclass
from pathlib import Path
from typing import List, Optional

from src.agent.schemas.document import (
    DocumentPage,
    DocumentSection,
    ParsedDocument,
)
from src.agent.parsers.base import BaseParser, ParseError

# Try to import python-docx
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


@dataclass
class WordParserConfig:
    """Configuration for Word parser."""
    extract_tables: bool = True
    detect_headings: bool = True


class WordParser(BaseParser):
    """Parser for Word (.docx) files.
    
    This parser:
    1. Extracts all paragraphs and their text
    2. Detects headings by paragraph style
    3. Extracts tables if enabled
    4. Preserves document structure
    
    Attributes:
        config: Parser configuration options
    """
    
    supported_extensions = [".docx"]
    
    # Heading style names in Word
    HEADING_STYLES = [
        'Heading 1', 'Heading 2', 'Heading 3', 'Heading 4',
        'Heading 5', 'Heading 6', 'Heading 1',
        '标题 1', '标题 2', '标题 3',  # Chinese
    ]
    
    def __init__(self, config: Optional[WordParserConfig] = None):
        """Initialize Word parser.
        
        Args:
            config: Optional configuration. Uses defaults if not provided.
        """
        self.config = config or WordParserConfig()
    
    def parse(self, file_path: str | Path) -> ParsedDocument:
        """Parse a Word document.
        
        Args:
            file_path: Path to the Word file.
            
        Returns:
            ParsedDocument with parsed content.
            
        Raises:
            ParseError: If parsing fails.
        """
        if not DOCX_AVAILABLE:
            raise ParseError(
                "python-docx is not installed. "
                "Install it with: pip install python-docx"
            )
        
        path = self.validate_file(file_path)
        
        try:
            doc = Document(str(path))
        except Exception as e:
            raise ParseError(f"Failed to open Word document: {e}") from e
        
        try:
            return self._parse_document(path, doc)
        except Exception as e:
            raise ParseError(f"Failed to parse Word document: {e}") from e
    
    def _parse_document(self, path: Path, doc: Document) -> ParsedDocument:
        """Parse Word document content.
        
        Args:
            path: File path
            doc: python-docx Document object
            
        Returns:
            ParsedDocument with pages and sections
        """
        sections: List[DocumentSection] = []
        tables_text: List[str] = []
        
        # Extract tables first if enabled
        if self.config.extract_tables:
            for table in doc.tables:
                table_text = self._extract_table_text(table)
                if table_text:
                    tables_text.append(table_text)
        
        # Process paragraphs
        paragraphs = list(doc.paragraphs)
        current_section_text = []
        current_section_title = None
        current_section_level = 1
        char_pos = 0
        
        for para in paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            # Check if paragraph is a heading
            is_heading, heading_level = self._is_heading(para)
            
            if is_heading:
                # Save previous section
                if current_section_text:
                    section_text = " ".join(current_section_text)
                    sections.append(DocumentSection(
                        title=current_section_title,
                        level=current_section_level,
                        page_no=1,
                        char_start=char_pos,
                        char_end=char_pos + len(section_text),
                        text=section_text,
                    ))
                    char_pos += len(section_text) + 1
                    current_section_text = []
                
                # Start new heading section
                current_section_title = text
                current_section_level = heading_level
            else:
                # Add to current section
                current_section_text.append(text)
        
        # Add final section
        if current_section_text:
            section_text = " ".join(current_section_text)
            sections.append(DocumentSection(
                title=current_section_title,
                level=current_section_level,
                page_no=1,
                char_start=char_pos,
                char_end=char_pos + len(section_text),
                text=section_text,
            ))
        
        # If no sections detected, create one from all text
        if not sections:
            all_text = "\n".join([p.text for p in paragraphs if p.text.strip()])
            sections.append(DocumentSection(
                title=None,
                level=1,
                page_no=1,
                char_start=0,
                char_end=len(all_text),
                text=all_text,
            ))
        
        # Add tables as sections if present
        if tables_text:
            for table_text in tables_text:
                sections.append(DocumentSection(
                    title="Table",
                    level=1,
                    page_no=1,
                    char_start=0,
                    char_end=len(table_text),
                    text=table_text,
                ))
        
        # Build full text
        full_text = "\n".join([s.text for s in sections])
        
        # Create single page (Word doesn't have pages in the same sense)
        page = DocumentPage(
            page_no=1,
            text=full_text,
            sections=sections,
        )
        
        return ParsedDocument(
            source_file=str(path),
            total_pages=1,
            pages=[page],
            sections=sections,
        )
    
    def _is_heading(self, para) -> tuple[bool, int]:
        """Check if paragraph is a heading.
        
        Args:
            para: python-docx Paragraph object
            
        Returns:
            Tuple of (is_heading, level)
        """
        if not self.config.detect_headings:
            return False, 1
        
        style_name = para.style.name if para.style else ""
        
        # Check for heading styles
        for heading_style in self.HEADING_STYLES:
            if heading_style in style_name:
                # Extract level from style name
                if 'Heading' in style_name:
                    try:
                        level = int(style_name.split()[-1])
                        return True, level
                    except (ValueError, IndexError):
                        pass
                elif '标题' in style_name:
                    try:
                        level = int(style_name[-1])
                        return True, level
                    except (ValueError, IndexError):
                        pass
                return True, 1
        
        # Fallback: short text with no ending punctuation might be heading
        text = para.text.strip()
        if text and len(text) < 80 and text[-1] not in '.?!;:':
            # Check if it looks like a title (title case)
            if text and text[0].isupper():
                words = text.split()
                if words and len(words) <= 10:
                    # Title case check: most words start with capital
                    title_case_count = sum(1 for w in words if w and w[0].isupper())
                    if title_case_count / len(words) > 0.5:
                        return True, 1
        
        return False, 1
    
    def _extract_table_text(self, table) -> str:
        """Extract text from a table.
        
        Args:
            table: python-docx Table object
            
        Returns:
            Table text as string
        """
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if cells:
                rows.append(" | ".join(cells))
        return "\n".join(rows)


__all__ = [
    "WordParser",
    "WordParserConfig",
]
