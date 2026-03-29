"""Word Document (DOCX) Loader implementation.

This module implements DOCX parsing for the ingestion pipeline,
converting Word documents to standardized Document format.

Features:
- Text extraction from paragraphs
- Heading detection by style
- TOC (Table of Contents) extraction
- Table extraction
- Structure preservation in Markdown format
"""

from __future__ import annotations

import hashlib
from pathlib import Path
from typing import Any, Dict, List, Optional

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from src.core.types import Document as CoreDocument
from src.libs.loader.base_loader import BaseLoader


class DocxLoader(BaseLoader):
    """DOCX Loader for Word documents.
    
    This loader:
    1. Extracts text from paragraphs
    2. Detects headings by paragraph style
    3. Extracts tables if present
    4. Converts to Markdown format
    5. Returns standardized Document object
    
    Configuration:
        extract_tables: Enable/disable table extraction (default: True)
        detect_headings: Enable/disable heading detection (default: True)
    
    Example:
        >>> loader = DocxLoader()
        >>> doc = loader.load("document.docx")
        >>> assert "source_path" in doc.metadata
    """
    
    HEADING_STYLES = [
        'Heading 1', 'Heading 2', 'Heading 3', 'Heading 4',
        'Heading 5', 'Heading 6',
        '标题 1', '标题 2', '标题 3',  # Chinese
    ]
    
    def __init__(
        self,
        extract_tables: bool = True,
        detect_headings: bool = True,
    ):
        """Initialize DOCX Loader.
        
        Args:
            extract_tables: Whether to extract tables from documents.
            detect_headings: Whether to detect headings by paragraph style.
        """
        if not DOCX_AVAILABLE:
            raise ImportError(
                "python-docx is required for DocxLoader. "
                "Install with: pip install python-docx"
            )
        
        self.extract_tables = extract_tables
        self.detect_headings = detect_headings
    
    def load(self, file_path: str | Path) -> CoreDocument:
        """Load and parse a Word document.
        
        Args:
            file_path: Path to the DOCX file.
            
        Returns:
            Document with Markdown text and metadata.
            
        Raises:
            FileNotFoundError: If the file doesn't exist.
            ValueError: If the file format is invalid.
            RuntimeError: If parsing fails critically.
        """
        # Validate file
        path = self._validate_file(file_path)
        if path.suffix.lower() not in ['.docx', '.doc']:
            raise ValueError(f"File is not a Word document: {path}")
        
        # Compute document hash for unique ID
        doc_hash = self._compute_file_hash(path)
        doc_id = f"doc_{doc_hash[:16]}"
        
        # Parse DOCX
        try:
            doc = Document(str(path))
        except Exception as e:
            raise RuntimeError(f"Failed to open Word document: {e}") from e
        
        # Extract TOC from document
        toc_entries = self._extract_toc(doc)
        
        # Extract content
        try:
            text_content, metadata = self._parse_document(path, doc, doc_hash)
        except Exception as e:
            raise RuntimeError(f"Failed to parse Word document: {e}") from e
        
        # Add TOC to metadata
        if toc_entries:
            metadata["toc"] = toc_entries
        
        return CoreDocument(
            id=doc_id,
            text=text_content,
            metadata=metadata
        )
    
    def _compute_file_hash(self, file_path: Path) -> str:
        """Compute SHA256 hash of file content."""
        sha256 = hashlib.sha256()
        with open(file_path, 'rb') as f:
            for chunk in iter(lambda: f.read(8192), b''):
                sha256.update(chunk)
        return sha256.hexdigest()
    
    def _parse_document(
        self,
        path: Path,
        doc,
        doc_hash: str
    ) -> tuple[str, Dict[str, Any]]:
        """Parse Word document content.
        
        Args:
            path: File path
            doc: python-docx Document object
            doc_hash: Document hash for ID generation
            
        Returns:
            Tuple of (markdown_text, metadata_dict)
        """
        sections: List[str] = []
        tables_text: List[str] = []
        title = None
        
        # Process paragraphs
        paragraphs = list(doc.paragraphs)
        
        for para in paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            # Check if paragraph is a heading
            is_heading, heading_level = self._is_heading(para)
            
            if is_heading:
                # Use first heading as document title
                if title is None:
                    title = text
                # Add as markdown heading
                heading_marker = '#' * min(heading_level, 6)
                sections.append(f"{heading_marker} {text}")
            else:
                # Add as regular paragraph
                sections.append(text)
        
        # Extract tables if enabled
        if self.extract_tables:
            for table in doc.tables:
                table_text = self._extract_table_markdown(table)
                if table_text:
                    tables_text.append(table_text)
                    sections.append("\n" + table_text + "\n")
        
        # Build full text
        text_content = "\n\n".join(sections)
        
        # If no content, try to get any text
        if not text_content.strip():
            text_content = "\n".join([p.text for p in paragraphs if p.text.strip()])
        
        # Initialize metadata
        metadata: Dict[str, Any] = {
            "source_path": str(path),
            "doc_type": "docx",
            "doc_hash": doc_hash,
            "paragraph_count": len([p for p in paragraphs if p.text.strip()]),
            "table_count": len(doc.tables),
        }
        
        if title:
            metadata["title"] = title
        
        return text_content, metadata
    
    def _is_heading(self, para) -> tuple[bool, int]:
        """Check if paragraph is a heading.
        
        Args:
            para: python-docx Paragraph object
            
        Returns:
            Tuple of (is_heading, level)
        """
        if not self.detect_headings:
            return False, 1
        
        style_name = para.style.name if para.style else ""
        
        # Check for heading styles
        for heading_style in self.HEADING_STYLES:
            if heading_style in style_name:
                # Extract level from style name
                if 'Heading' in style_name or '标题' in style_name:
                    try:
                        # Try to extract number at end
                        level = int(style_name[-1])
                        return True, level
                    except (ValueError, IndexError):
                        pass
                return True, 1
        
        # Fallback: short text without ending punctuation might be heading
        text = para.text.strip()
        if text and len(text) < 80 and text[-1] not in '.?!;:':
            if text and text[0].isupper():
                words = text.split()
                if words and len(words) <= 10:
                    title_case_count = sum(1 for w in words if w and w[0].isupper())
                    if title_case_count / len(words) > 0.5:
                        return True, 1
        
        return False, 1
    
    def _extract_toc(self, doc) -> List[Dict[str, Any]]:
        """Extract Table of Contents from Word document.
        
        DOCX TOC extraction requires parsing the document's XML structure
        to find TOC fields. This method extracts headings which typically
        represent the TOC structure.
        
        Args:
            doc: python-docx Document object
            
        Returns:
            List of TOC entries with title, level, and position.
        """
        toc_entries = []
        
        try:
            paragraphs = list(doc.paragraphs)
            
            for para in paragraphs:
                text = para.text.strip()
                if not text:
                    continue
                
                # Check if paragraph is a heading (TOC entries are usually headings)
                is_heading, heading_level = self._is_heading(para)
                
                if is_heading:
                    toc_entries.append({
                        "title": text,
                        "level": heading_level,
                        "type": "heading"
                    })
                else:
                    # Check for common TOC patterns in non-heading text
                    if self._is_toc_entry(text):
                        # Estimate level based on indentation or pattern
                        level = self._estimate_toc_level(para)
                        toc_entries.append({
                            "title": text,
                            "level": level,
                            "type": "toc_entry"
                        })
            
            return toc_entries
            
        except Exception as e:
            import logging
            logger = logging.getLogger(__name__)
            logger.warning(f"Failed to extract TOC from DOCX: {e}")
            return []
    
    def _is_toc_entry(self, text: str) -> bool:
        """Check if text matches TOC entry patterns.
        
        Args:
            text: Text to check
            
        Returns:
            True if text matches TOC entry patterns.
        """
        import re
        
        # TOC entry patterns (Chinese and English)
        toc_patterns = [
            r'^第[一二三四五六七八九十百千\d]+[章节篇部节]',  # 第一章, 第一节
            r'^\d+\.\d+',  # 1.1, 2.3
            r'^[一二三四五六七八九十]+、',  # 一、, 二、
            r'^(Chapter|Section|Article)\s+\d+',  # Chapter 1, Section 2
            r'^(目录|Contents|Table of Contents)',  # TOC header
        ]
        
        for pattern in toc_patterns:
            if re.search(pattern, text):
                return True
        
        return False
    
    def _estimate_toc_level(self, para) -> int:
        """Estimate TOC entry level based on paragraph properties.
        
        Args:
            para: python-docx Paragraph object
            
        Returns:
            Estimated heading level (1-6).
        """
        # Check for outline level (if available in XML)
        try:
            # Try to get outline level from paragraph format
            para_xml = para._element
            pPr = para_xml.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr')
            if pPr is not None:
                pStyle = pPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pStyle')
                if pStyle is not None:
                    style_val = pStyle.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', '')
                    if 'Heading' in style_val or '标题' in style_val:
                        # Extract level from style name
                        import re
                        match = re.search(r'(\d+)$', style_val)
                        if match:
                            return int(match.group(1))
        except Exception:
            pass
        
        # Default to level 1 for unrecognized entries
        return 1
    
    def _extract_table_markdown(self, table) -> str:
        """Extract table as Markdown.
        
        Args:
            table: python-docx Table object
            
        Returns:
            Table text as Markdown string
        """
        rows_data = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if cells:
                rows_data.append(cells)
        
        if not rows_data:
            return ""
        
        # Build markdown table
        lines = []
        for i, row in enumerate(rows_data):
            lines.append("| " + " | ".join(row) + " |")
            # Add separator after header row
            if i == 0:
                separator = "| " + " | ".join(["---"] * len(row)) + " |"
                lines.append(separator)
        
        return "\n".join(lines)
